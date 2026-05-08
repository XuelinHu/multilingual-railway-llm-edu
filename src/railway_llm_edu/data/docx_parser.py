from __future__ import annotations

from pathlib import Path

from docx import Document

from railway_llm_edu.schemas import RawParagraph
from railway_llm_edu.utils import normalize_space


def parse_docx(path: str | Path, doc_id: str, min_chars: int = 2) -> list[RawParagraph]:
    """Extract paragraphs and table cells from a docx file in reading order."""
    path = Path(path)
    document = Document(path)
    rows: list[RawParagraph] = []

    for idx, para in enumerate(document.paragraphs):
        text = normalize_space(para.text)
        if len(text) >= min_chars:
            rows.append(
                RawParagraph(
                    doc_id=doc_id,
                    source_path=str(path),
                    paragraph_id=f"{doc_id}-p{idx:05d}",
                    text=text,
                    style=para.style.name if para.style else None,
                )
            )

    for t_idx, table in enumerate(document.tables):
        for r_idx, row in enumerate(table.rows):
            cells = [normalize_space(cell.text) for cell in row.cells]
            text = "\t".join(cell for cell in cells if cell)
            if len(text) >= min_chars:
                rows.append(
                    RawParagraph(
                        doc_id=doc_id,
                        source_path=str(path),
                        paragraph_id=f"{doc_id}-t{t_idx:03d}-r{r_idx:05d}",
                        text=text,
                        table_id=f"{doc_id}-t{t_idx:03d}",
                        row_id=r_idx,
                    )
                )
    return rows
